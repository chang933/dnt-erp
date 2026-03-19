# -*- coding: utf-8 -*-
"""
직원 서류 양식 생성기
6종 서류를 docx 형식으로 생성합니다.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
import os
from typing import Optional

# 이미지 경로 설정
IMAGES_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "images")


def add_watermark_to_section(section):
    """페이지 헤더에 워터마크 로고 추가"""
    try:
        logo_path = os.path.join(IMAGES_DIR, "graylogo.png")
        if os.path.exists(logo_path):
            # 헤더에 워터마크 이미지 추가
            header = section.header
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 기존 내용 제거
            header_para.clear()
            
            # 워터마크 이미지 추가 (floating으로 중앙 배치)
            run = header_para.add_run()
            try:
                run.add_picture(logo_path, width=Cm(7))  # 약 7cm 크기
                # 이미지를 투명하게 만들기 위해 XML 수정 필요 (복잡하므로 일단 추가만)
            except Exception as e:
                print(f"워터마크 이미지 추가 실패: {e}")
    except Exception as e:
        print(f"워터마크 추가 중 오류: {e}")


def add_stamp_image(paragraph):
    """서명 부분에 직인 이미지 추가"""
    try:
        stamp_path = os.path.join(IMAGES_DIR, "stamp.png")
        if os.path.exists(stamp_path):
            # paragraph에 직인 이미지 추가
            run = paragraph.add_run()
            try:
                run.add_picture(stamp_path, width=Cm(2.5))  # 약 2.5cm 크기
            except Exception as e:
                print(f"직인 이미지 추가 실패: {e}")
    except Exception as e:
        print(f"직인 추가 중 오류: {e}")


def _get_base64_image(filename: str) -> str:
    """이미지를 base64로 인코딩"""
    try:
        image_path = os.path.join(IMAGES_DIR, filename)
        if os.path.exists(image_path):
            import base64
            with open(image_path, 'rb') as f:
                return base64.b64encode(f.read()).decode('utf-8')
    except Exception as e:
        print(f"이미지 로드 실패 {filename}: {e}")
    return ''


def _format_position_title(employee_data: dict) -> str:
    """직위 포맷팅 (급여 형태와 포지션에 따라)"""
    position = employee_data.get('employee_position', '')
    salary_type = employee_data.get('salary_type', '')
    
    # 대표, 사장은 그대로 표시
    if position in ['대표', '사장']:
        return position
    
    # 홀/주방 + 월급/시급 조합
    if salary_type == '월급':
        if position == '홀':
            return '홀 직원'
        elif position == '주방':
            return '주방 직원'
    elif salary_type == '시급':
        if position == '홀':
            return '홀 아르바이트'
        elif position == '주방':
            return '주방 아르바이트'
    
    # 기본값: 원래 포지션 그대로
    return position


def _format_job_description(employee_data: dict) -> str:
    """담당업무 포맷팅"""
    position = employee_data.get('employee_position', '')
    
    # 대표, 사장은 그대로 표시
    if position in ['대표', '사장']:
        return position
    
    # 홀/주방 처리
    if position == '홀':
        return '홀 서빙'
    elif position == '주방':
        return '주방'
    
    # 기본값
    return position


def create_receipt_of_employment_html(employee_data: dict) -> str:
    """재직증명서 생성 (HTML 형식)"""
    today = datetime.now().strftime('%Y년 %m월 %d일')
    
    html = f"""
    <!DOCTYPE html>
    <html lang="ko">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>재직증명서</title>
        <style>
            @page {{
                size: A4;
                margin: 2cm;
            }}
            * {{
                margin: 0;
                padding: 0;
                box-sizing: border-box;
            }}
            body {{
                font-family: '맑은 고딕', 'Malgun Gothic', sans-serif;
                font-size: 11pt;
                line-height: 1.6;
                color: #000;
                padding: 40px;
                background: #fff;
            }}
            .container {{
                max-width: 21cm;
                margin: 0 auto;
                background: white;
                position: relative;
            }}
            .watermark {{
                position: absolute;
                top: 50%;
                left: 50%;
                transform: translate(-50%, -50%);
                opacity: 0.1;
                z-index: 0;
                pointer-events: none;
            }}
            .watermark img {{
                width: 280px;
                height: 280px;
            }}
            .content {{
                position: relative;
                z-index: 1;
            }}
            .title {{
                text-align: center;
                font-size: 24pt;
                font-weight: bold;
                margin-bottom: 10px;
                letter-spacing: 2px;
            }}
            .title-underline {{
                text-align: center;
                border-bottom: 2px solid #000;
                margin: 0 auto 70px;
                width: 280px;
            }}
            .body-text {{
                text-align: center;
                font-size: 12pt;
                margin-bottom: 30px;
            }}
            table {{
                width: 100%;
                border-collapse: collapse;
                margin-bottom: 50px;
            }}
            table td {{
                border: 1px solid #000;
                padding: 18px;
                vertical-align: middle;
                height: 1.5em;
            }}
            table td.label {{
                width: 25%;
                text-align: center;
                font-weight: bold;
                background-color: #f9f9f9;
            }}
            table td.value {{
                width: 75%;
                text-align: left;
                padding-left: 15px;
            }}
            .signature {{
                text-align: center;
                margin-top: 50px;
            }}
            .date {{
                font-size: 13pt;
                margin-bottom: 40px;
            }}
            .company-signature {{
                font-size: 18pt;
                font-weight: bold;
                margin-top: 10px;
                position: relative;
                display: inline-block;
            }}
            .ceo-text {{
                display: inline-block;
                font-size: 13pt;
                margin-left: 10px;
                position: relative;
            }}
            .stamp {{
                position: absolute;
                right: -5px;
                top: -15px;
                z-index: 10;
            }}
            .stamp img {{
                width: 50px;
                height: 50px;
            }}
            .in-text {{
                position: relative;
                display: inline-block;
            }}
            @media print {{
                body {{
                    padding: 0;
                }}
                .container {{
                    max-width: 100%;
                }}
            }}
        </style>
    </head>
    <body>
        <div class="container">
            <div class="watermark">
                <img src="data:image/png;base64,{_get_base64_image('graylogo.png')}" alt="워터마크">
            </div>
            <div class="content">
                <div class="title">재&nbsp;&nbsp;직&nbsp;&nbsp;증&nbsp;&nbsp;명&nbsp;&nbsp;서</div>
                <div class="title-underline"></div>
                <br>
                <table>
                    <tr>
                        <td class="label">성&nbsp;&nbsp;&nbsp;&nbsp;명</td>
                        <td class="value">{employee_data.get('name', '')}</td>
                    </tr>
                    <tr>
                        <td class="label">주민등록번호</td>
                        <td class="value">{employee_data.get('ssn', '') or ''}</td>
                    </tr>
                    <tr>
                        <td class="label">생년월일</td>
                        <td class="value">{employee_data.get('birth_date', '') or ''}</td>
                    </tr>
                    <tr>
                        <td class="label">주&nbsp;&nbsp;&nbsp;&nbsp;소</td>
                        <td class="value">{employee_data.get('address', '') or ''}</td>
                    </tr>
                    <tr>
                        <td class="label">소&nbsp;&nbsp;&nbsp;&nbsp;속</td>
                        <td class="value">도원반점 검단점</td>
                    </tr>
                    <tr>
                        <td class="label">직&nbsp;&nbsp;&nbsp;&nbsp;위</td>
                        <td class="value">{_format_position_title(employee_data)}</td>
                    </tr>
                    <tr>
                        <td class="label">입 사 일</td>
                        <td class="value">{employee_data.get('hire_date', '') or ''}</td>
                    </tr>
                    <tr>
                        <td class="label">퇴 사 일</td>
                        <td class="value">{employee_data.get('resign_date', '') or ''}</td>
                    </tr>
                </table>
                <br>
                <div class="body-text">위 사람은 당 업체에 재직 중임을 증명합니다.</div>
                <div class="signature">
                    <div class="date">{today}</div>
                    <div class="company-signature">
                        도원반점 검단점
                        <span class="ceo-text">사장 김서은 (<span class="in-text">인<span class="stamp">
                            <img src="data:image/png;base64,{_get_base64_image('stamp.png')}" alt="직인">
                        </span></span>)</span>
                    </div>
                </div>
            </div>
        </div>
    </body>
    </html>
    """
    return html


def create_career_certificate_html(employee_data: dict) -> str:
    """경력증명서 생성 (HTML 형식)"""
    today = datetime.now().strftime('%Y년 %m월 %d일')
    
    # 경력 기간 계산
    hire_date = employee_data.get('hire_date', '')
    resign_date = employee_data.get('resign_date', '')
    career_period = ''
    if hire_date:
        from datetime import datetime as dt
        try:
            hire = dt.strptime(hire_date, '%Y-%m-%d')
            if resign_date:
                resign = dt.strptime(resign_date, '%Y-%m-%d')
            else:
                resign = dt.now()
            years = resign.year - hire.year
            months = resign.month - hire.month
            if months < 0:
                years -= 1
                months += 12
            career_period = f'{years}년 {months}개월'
        except:
            career_period = ''
    
    html = f"""
    <!DOCTYPE html>
    <html lang="ko">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>경력증명서</title>
        <style>
            @page {{
                size: A4;
                margin: 2cm;
            }}
            * {{
                margin: 0;
                padding: 0;
                box-sizing: border-box;
            }}
            body {{
                font-family: '맑은 고딕', 'Malgun Gothic', sans-serif;
                font-size: 11pt;
                line-height: 1.6;
                color: #000;
                padding: 40px;
                background: #fff;
            }}
            .container {{
                max-width: 21cm;
                margin: 0 auto;
                background: white;
                position: relative;
            }}
            .watermark {{
                position: absolute;
                top: 50%;
                left: 50%;
                transform: translate(-50%, -50%);
                opacity: 0.1;
                z-index: 0;
                pointer-events: none;
            }}
            .watermark img {{
                width: 280px;
                height: 280px;
            }}
            .content {{
                position: relative;
                z-index: 1;
            }}
            .title {{
                text-align: center;
                font-size: 24pt;
                font-weight: bold;
                margin-bottom: 10px;
                letter-spacing: 2px;
            }}
            .title-underline {{
                text-align: center;
                border-bottom: 2px solid #000;
                margin: 0 auto 70px;
                width: 280px;
            }}
            .body-text {{
                text-align: center;
                font-size: 12pt;
                margin-bottom: 30px;
            }}
            table {{
                width: 100%;
                border-collapse: collapse;
                margin-bottom: 50px;
            }}
            table td {{
                border: 1px solid #000;
                padding: 18px;
                vertical-align: middle;
                height: 1.5em;
            }}
            table td.label {{
                width: 25%;
                text-align: center;
                font-weight: bold;
                background-color: #f9f9f9;
            }}
            table td.value {{
                width: 75%;
                text-align: left;
                padding-left: 15px;
            }}
            .signature {{
                text-align: center;
                margin-top: 50px;
            }}
            .date {{
                font-size: 13pt;
                margin-bottom: 40px;
            }}
            .company-signature {{
                font-size: 18pt;
                font-weight: bold;
                margin-top: 10px;
                position: relative;
                display: inline-block;
            }}
            .ceo-text {{
                display: inline-block;
                font-size: 13pt;
                margin-left: 10px;
                position: relative;
            }}
            .stamp {{
                position: absolute;
                right: -5px;
                top: -15px;
                z-index: 10;
            }}
            .stamp img {{
                width: 50px;
                height: 50px;
            }}
            .in-text {{
                position: relative;
                display: inline-block;
            }}
            @media print {{
                body {{
                    padding: 0;
                }}
                .container {{
                    max-width: 100%;
                }}
            }}
        </style>
    </head>
    <body>
        <div class="container">
            <div class="watermark">
                <img src="data:image/png;base64,{_get_base64_image('graylogo.png')}" alt="워터마크">
            </div>
            <div class="content">
                <div class="title">경&nbsp;&nbsp;력&nbsp;&nbsp;증&nbsp;&nbsp;명&nbsp;&nbsp;서</div>
                <div class="title-underline"></div>
                <br>
                
                <table>
                    <tr>
                        <td class="label">성&nbsp;&nbsp;&nbsp;&nbsp;명</td>
                        <td class="value">{employee_data.get('name', '')}</td>
                    </tr>
                    <tr>
                        <td class="label">주민등록번호</td>
                        <td class="value">{employee_data.get('ssn', '') or ''}</td>
                    </tr>
                    <tr>
                        <td class="label">생년월일</td>
                        <td class="value">{employee_data.get('birth_date', '') or ''}</td>
                    </tr>
                    <tr>
                        <td class="label">주&nbsp;&nbsp;&nbsp;&nbsp;소</td>
                        <td class="value">{employee_data.get('address', '') or ''}</td>
                    </tr>
                    <tr>
                        <td class="label">소&nbsp;&nbsp;&nbsp;&nbsp;속</td>
                        <td class="value">도원반점 검단점</td>
                    </tr>
                    <tr>
                        <td class="label">직&nbsp;&nbsp;&nbsp;&nbsp;위</td>
                        <td class="value">{_format_position_title(employee_data)}</td>
                    </tr>
                    <tr>
                        <td class="label">입 사 일</td>
                        <td class="value">{employee_data.get('hire_date', '') or ''}</td>
                    </tr>
                    <tr>
                        <td class="label">퇴 사 일</td>
                        <td class="value">{employee_data.get('resign_date', '') or '재직중'}</td>
                    </tr>
                    <tr>
                        <td class="label">담당업무</td>
                        <td class="value">{_format_job_description(employee_data)}</td>
                    </tr>
                    <tr>
                        <td class="label">경&nbsp;&nbsp;&nbsp;&nbsp;력</td>
                        <td class="value">{career_period}</td>
                    </tr>
                </table>
                <div class="body-text">위 사람은 본 식당에 근무한 경력이 아래와 같음을 증명합니다.</div>
                <div class="signature">
                    <div class="date">{today}</div>
                    <div class="company-signature">
                        도원반점 검단점
                        <span class="ceo-text">대표 김서은 (<span class="in-text">인<span class="stamp">
                            <img src="data:image/png;base64,{_get_base64_image('stamp.png')}" alt="직인">
                        </span></span>)</span>
                    </div>
                </div>
            </div>
        </div>
    </body>
    </html>
    """
    return html


def create_pay_stub_html(employee_data: dict, payroll_data: Optional[dict] = None) -> str:
    """급여명세서 생성 (HTML 형식)"""
    today = datetime.now().strftime('%Y년 %m월 %d일')
    
    payroll_year_month = ''
    base_pay = 0
    deductions = 0
    net_pay = 0
    
    if payroll_data:
        payroll_year_month = payroll_data.get('year_month', '')
        base_pay = int(payroll_data.get('base_pay', 0))
        deductions = int(payroll_data.get('deductions', 0))
        net_pay = int(payroll_data.get('net_pay', 0))
    else:
        from datetime import datetime as dt
        now = dt.now()
        payroll_year_month = f'{now.year}-{now.month:02d}'
    
    html = f"""
    <!DOCTYPE html>
    <html lang="ko">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>급여명세서</title>
        <style>
            @page {{
                size: A4;
                margin: 2cm;
            }}
            * {{
                margin: 0;
                padding: 0;
                box-sizing: border-box;
            }}
            body {{
                font-family: '맑은 고딕', 'Malgun Gothic', sans-serif;
                font-size: 11pt;
                line-height: 1.6;
                color: #000;
                padding: 40px;
                background: #fff;
            }}
            .container {{
                max-width: 21cm;
                margin: 0 auto;
                background: white;
                position: relative;
            }}
            .watermark {{
                position: absolute;
                top: 50%;
                left: 50%;
                transform: translate(-50%, -50%);
                opacity: 0.1;
                z-index: 0;
                pointer-events: none;
            }}
            .watermark img {{
                width: 280px;
                height: 280px;
            }}
            .content {{
                position: relative;
                z-index: 1;
            }}
            .title {{
                text-align: center;
                font-size: 24pt;
                font-weight: bold;
                margin-bottom: 10px;
                letter-spacing: 2px;
            }}
            .title-underline {{
                text-align: center;
                border-bottom: 2px solid #000;
                margin: 0 auto 70px;
                width: 280px;
            }}
            table {{
                width: 100%;
                border-collapse: collapse;
                margin-bottom: 50px;
            }}
            table td {{
                border: 1px solid #000;
                padding: 18px;
                vertical-align: middle;
                height: 1.5em;
            }}
            table td.label {{
                width: 25%;
                text-align: center;
                font-weight: bold;
                background-color: #f9f9f9;
            }}
            table td.value {{
                width: 75%;
                text-align: left;
                padding-left: 15px;
            }}
            .signature {{
                text-align: center;
                margin-top: 50px;
            }}
            .date {{
                font-size: 13pt;
                margin-bottom: 40px;
            }}
            .company-signature {{
                font-size: 18pt;
                font-weight: bold;
                margin-top: 10px;
                position: relative;
                display: inline-block;
            }}
            .ceo-text {{
                display: inline-block;
                font-size: 13pt;
                margin-left: 10px;
                position: relative;
            }}
            .stamp {{
                position: absolute;
                right: -5px;
                top: -15px;
                z-index: 10;
            }}
            .stamp img {{
                width: 50px;
                height: 50px;
            }}
            .in-text {{
                position: relative;
                display: inline-block;
            }}
            @media print {{
                body {{
                    padding: 0;
                }}
                .container {{
                    max-width: 100%;
                }}
            }}
        </style>
    </head>
    <body>
        <div class="container">
            <div class="watermark">
                <img src="data:image/png;base64,{_get_base64_image('graylogo.png')}" alt="워터마크">
            </div>
            <div class="content">
                <div class="title">급&nbsp;&nbsp;여&nbsp;&nbsp;명&nbsp;&nbsp;세&nbsp;&nbsp;서</div>
                <div class="title-underline"></div>
                <br><br><br>
                <table>
                    <tr>
                        <td class="label">성&nbsp;&nbsp;&nbsp;&nbsp;명</td>
                        <td class="value">{employee_data.get('name', '')}</td>
                    </tr>
                    <tr>
                        <td class="label">주민등록번호</td>
                        <td class="value">{employee_data.get('ssn', '') or ''}</td>
                    </tr>
                    <tr>
                        <td class="label">소&nbsp;&nbsp;&nbsp;&nbsp;속</td>
                        <td class="value">도원반점 검단점</td>
                    </tr>
                    <tr>
                        <td class="label">지 급 년 월</td>
                        <td class="value">{payroll_year_month}</td>
                    </tr>
                </table>
                <table>
                    <tr>
                        <td class="label">기본급</td>
                        <td class="value">{base_pay:,}원</td>
                    </tr>
                    <tr>
                        <td class="label">공제액</td>
                        <td class="value">{deductions:,}원</td>
                    </tr>
                    <tr>
                        <td class="label">실수령액</td>
                        <td class="value">{net_pay:,}원</td>
                    </tr>
                </table>
                <div class="signature">
                    <div class="date">{today}</div>
                    <div class="company-signature">
                        도원반점 검단점
                        <span class="ceo-text">사장 김서은 (<span class="in-text">인<span class="stamp">
                            <img src="data:image/png;base64,{_get_base64_image('stamp.png')}" alt="직인">
                        </span></span>)</span>
                    </div>
                </div>
            </div>
        </div>
    </body>
    </html>
    """
    return html


def create_withholding_receipt_html(employee_data: dict, payroll_data: Optional[dict] = None) -> str:
    """원천징수영수증 생성 (HTML 형식)"""
    today = datetime.now().strftime('%Y년 %m월 %d일')
    
    payroll_year_month = ''
    base_pay = 0
    deductions = 0
    
    if payroll_data:
        payroll_year_month = payroll_data.get('year_month', '')
        base_pay = int(payroll_data.get('base_pay', 0))
        deductions = int(payroll_data.get('deductions', 0))
    else:
        from datetime import datetime as dt
        now = dt.now()
        payroll_year_month = f'{now.year}-{now.month:02d}'
    
    html = f"""
    <!DOCTYPE html>
    <html lang="ko">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>원천징수영수증</title>
        <style>
            @page {{
                size: A4;
                margin: 2cm;
            }}
            * {{
                margin: 0;
                padding: 0;
                box-sizing: border-box;
            }}
            body {{
                font-family: '맑은 고딕', 'Malgun Gothic', sans-serif;
                font-size: 11pt;
                line-height: 1.6;
                color: #000;
                padding: 40px;
                background: #fff;
            }}
            .container {{
                max-width: 21cm;
                margin: 0 auto;
                background: white;
                position: relative;
            }}
            .watermark {{
                position: absolute;
                top: 50%;
                left: 50%;
                transform: translate(-50%, -50%);
                opacity: 0.1;
                z-index: 0;
                pointer-events: none;
            }}
            .watermark img {{
                width: 280px;
                height: 280px;
            }}
            .content {{
                position: relative;
                z-index: 1;
            }}
            .title {{
                text-align: center;
                font-size: 24pt;
                font-weight: bold;
                margin-bottom: 10px;
                letter-spacing: 2px;
            }}
            .title-underline {{
                text-align: center;
                border-bottom: 2px solid #000;
                margin: 0 auto 70px;
                width: 280px;
            }}
            table {{
                width: 100%;
                border-collapse: collapse;
                margin-bottom: 50px;
            }}
            table td {{
                border: 1px solid #000;
                padding: 18px;
                vertical-align: middle;
                height: 1.5em;
            }}
            table td.label {{
                width: 25%;
                text-align: center;
                font-weight: bold;
                background-color: #f9f9f9;
            }}
            table td.value {{
                width: 75%;
                text-align: left;
                padding-left: 15px;
            }}
            .signature {{
                text-align: center;
                margin-top: 50px;
            }}
            .date {{
                font-size: 13pt;
                margin-bottom: 40px;
            }}
            .company-signature {{
                font-size: 18pt;
                font-weight: bold;
                margin-top: 10px;
                position: relative;
                display: inline-block;
            }}
            .ceo-text {{
                display: inline-block;
                font-size: 13pt;
                margin-left: 10px;
                position: relative;
            }}
            .stamp {{
                position: absolute;
                right: -5px;
                top: -15px;
                z-index: 10;
            }}
            .stamp img {{
                width: 50px;
                height: 50px;
            }}
            .in-text {{
                position: relative;
                display: inline-block;
            }}
            @media print {{
                body {{
                    padding: 0;
                }}
                .container {{
                    max-width: 100%;
                }}
            }}
        </style>
    </head>
    <body>
        <div class="container">
            <div class="watermark">
                <img src="data:image/png;base64,{_get_base64_image('graylogo.png')}" alt="워터마크">
            </div>
            <div class="content">
                <div class="title">원&nbsp;천&nbsp;징&nbsp;수&nbsp;영&nbsp;수&nbsp;증</div>
                <div class="title-underline"></div>
                <br><br>
                <table>
                    <tr>
                        <td class="label">성&nbsp;&nbsp;&nbsp;&nbsp;명</td>
                        <td class="value">{employee_data.get('name', '')}</td>
                    </tr>
                    <tr>
                        <td class="label">주민등록번호</td>
                        <td class="value">{employee_data.get('ssn', '') or ''}</td>
                    </tr>
                    <tr>
                        <td class="label">소&nbsp;&nbsp;&nbsp;&nbsp;속</td>
                        <td class="value">도원반점 검단점</td>
                    </tr>
                    <tr>
                        <td class="label">지 급 년 월</td>
                        <td class="value">{payroll_year_month}</td>
                    </tr>
                </table>
                <table>
                    <tr>
                        <td class="label">지 급 금 액</td>
                        <td class="value">{base_pay:,}원</td>
                    </tr>
                    <tr>
                        <td class="label">원천징수세액</td>
                        <td class="value">{deductions:,}원</td>
                    </tr>
                </table>
                <div class="signature">
                    <div class="date">{today}</div>
                    <div class="company-signature">
                        도원반점 검단점
                        <span class="ceo-text">사장 김서은 (<span class="in-text">인<span class="stamp">
                            <img src="data:image/png;base64,{_get_base64_image('stamp.png')}" alt="직인">
                        </span></span>)</span>
                    </div>
                </div>
            </div>
        </div>
    </body>
    </html>
    """
    return html


def create_resignation_certificate_html(employee_data: dict) -> str:
    """퇴직증명서 생성 (HTML 형식)"""
    today = datetime.now().strftime('%Y년 %m월 %d일')
    
    resign_date = employee_data.get('resign_date', '')
    if not resign_date:
        resign_date = datetime.now().strftime('%Y-%m-%d')
    
    html = f"""
    <!DOCTYPE html>
    <html lang="ko">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>퇴직증명서</title>
        <style>
            @page {{
                size: A4;
                margin: 2cm;
            }}
            * {{
                margin: 0;
                padding: 0;
                box-sizing: border-box;
            }}
            body {{
                font-family: '맑은 고딕', 'Malgun Gothic', sans-serif;
                font-size: 11pt;
                line-height: 1.6;
                color: #000;
                padding: 40px;
                background: #fff;
            }}
            .container {{
                max-width: 21cm;
                margin: 0 auto;
                background: white;
                position: relative;
            }}
            .watermark {{
                position: absolute;
                top: 50%;
                left: 50%;
                transform: translate(-50%, -50%);
                opacity: 0.1;
                z-index: 0;
                pointer-events: none;
            }}
            .watermark img {{
                width: 280px;
                height: 280px;
            }}
            .content {{
                position: relative;
                z-index: 1;
            }}
            .title {{
                text-align: center;
                font-size: 24pt;
                font-weight: bold;
                margin-bottom: 10px;
                letter-spacing: 2px;
            }}
            .title-underline {{
                text-align: center;
                border-bottom: 2px solid #000;
                margin: 0 auto 70px;
                width: 280px;
            }}
            .body-text {{
                text-align: center;
                font-size: 12pt;
                margin-bottom: 30px;
            }}
            table {{
                width: 100%;
                border-collapse: collapse;
                margin-bottom: 50px;
            }}
            table td {{
                border: 1px solid #000;
                padding: 18px;
                vertical-align: middle;
                height: 1.5em;
            }}
            table td.label {{
                width: 25%;
                text-align: center;
                font-weight: bold;
                background-color: #f9f9f9;
            }}
            table td.value {{
                width: 75%;
                text-align: left;
                padding-left: 15px;
            }}
            .signature {{
                text-align: center;
                margin-top: 50px;
            }}
            .date {{
                font-size: 13pt;
                margin-bottom: 40px;
            }}
            .company-signature {{
                font-size: 18pt;
                font-weight: bold;
                margin-top: 10px;
                position: relative;
                display: inline-block;
            }}
            .ceo-text {{
                display: inline-block;
                font-size: 13pt;
                margin-left: 10px;
                position: relative;
            }}
            .stamp {{
                position: absolute;
                right: -5px;
                top: -15px;
                z-index: 10;
            }}
            .stamp img {{
                width: 50px;
                height: 50px;
            }}
            .in-text {{
                position: relative;
                display: inline-block;
            }}
            @media print {{
                body {{
                    padding: 0;
                }}
                .container {{
                    max-width: 100%;
                }}
            }}
        </style>
    </head>
    <body>
        <div class="container">
            <div class="watermark">
                <img src="data:image/png;base64,{_get_base64_image('graylogo.png')}" alt="워터마크">
            </div>
            <div class="content">
                <div class="title">퇴&nbsp;&nbsp;직&nbsp;&nbsp;증&nbsp;&nbsp;명&nbsp;&nbsp;서</div>
                <div class="title-underline"></div>
                <br><br>
                
                <table>
                    <tr>
                        <td class="label">성&nbsp;&nbsp;&nbsp;&nbsp;명</td>
                        <td class="value">{employee_data.get('name', '')}</td>
                    </tr>
                    <tr>
                        <td class="label">주민등록번호</td>
                        <td class="value">{employee_data.get('ssn', '') or ''}</td>
                    </tr>
                    <tr>
                        <td class="label">생년월일</td>
                        <td class="value">{employee_data.get('birth_date', '') or ''}</td>
                    </tr>
                    <tr>
                        <td class="label">주&nbsp;&nbsp;&nbsp;&nbsp;소</td>
                        <td class="value">{employee_data.get('address', '') or ''}</td>
                    </tr>
                    <tr>
                        <td class="label">소&nbsp;&nbsp;&nbsp;&nbsp;속</td>
                        <td class="value">도원반점 검단점</td>
                    </tr>
                    <tr>
                        <td class="label">직&nbsp;&nbsp;&nbsp;&nbsp;위</td>
                        <td class="value">{_format_position_title(employee_data)}</td>
                    </tr>
                    <tr>
                        <td class="label">입 사 일</td>
                        <td class="value">{employee_data.get('hire_date', '') or ''}</td>
                    </tr>
                    <tr>
                        <td class="label">퇴 사 일</td>
                        <td class="value">{resign_date}</td>
                    </tr>
                    <tr>
                        <td class="label">담당업무</td>
                        <td class="value">{_format_job_description(employee_data)}</td>
                    </tr>
                </table>
                <div class="body-text">위 사람은 본 식당의 직원으로 재직하다가 아래와 같이 퇴직함을 증명합니다.</div>
                <div class="signature">
                    <div class="date">{today}</div>
                    <div class="company-signature">
                        도원반점 검단점
                        <span class="ceo-text">사장 김서은 (<span class="in-text">인<span class="stamp">
                            <img src="data:image/png;base64,{_get_base64_image('stamp.png')}" alt="직인">
                        </span></span>)</span>
                    </div>
                </div>
            </div>
        </div>
    </body>
    </html>
    """
    return html


def create_severance_settlement_html(employee_data: dict, severance_data: Optional[dict] = None) -> str:
    """퇴직금정산서 생성 (HTML 형식)"""
    today = datetime.now().strftime('%Y년 %m월 %d일')
    
    resign_date = employee_data.get('resign_date', '')
    if not resign_date:
        resign_date = datetime.now().strftime('%Y-%m-%d')
    
    severance_amount = 0
    if severance_data:
        severance_amount = int(severance_data.get('severance_amount', 0))
    
    html = f"""
    <!DOCTYPE html>
    <html lang="ko">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>퇴직금정산서</title>
        <style>
            @page {{
                size: A4;
                margin: 2cm;
            }}
            * {{
                margin: 0;
                padding: 0;
                box-sizing: border-box;
            }}
            body {{
                font-family: '맑은 고딕', 'Malgun Gothic', sans-serif;
                font-size: 11pt;
                line-height: 1.6;
                color: #000;
                padding: 40px;
                background: #fff;
            }}
            .container {{
                max-width: 21cm;
                margin: 0 auto;
                background: white;
                position: relative;
            }}
            .watermark {{
                position: absolute;
                top: 50%;
                left: 50%;
                transform: translate(-50%, -50%);
                opacity: 0.1;
                z-index: 0;
                pointer-events: none;
            }}
            .watermark img {{
                width: 280px;
                height: 280px;
            }}
            .content {{
                position: relative;
                z-index: 1;
            }}
            .title {{
                text-align: center;
                font-size: 24pt;
                font-weight: bold;
                margin-bottom: 10px;
                letter-spacing: 2px;
            }}
            .title-underline {{
                text-align: center;
                border-bottom: 2px solid #000;
                margin: 0 auto 70px;
                width: 280px;
            }}
            table {{
                width: 100%;
                border-collapse: collapse;
                margin-bottom: 50px;
            }}
            table td {{
                border: 1px solid #000;
                padding: 18px;
                vertical-align: middle;
                height: 1.5em;
            }}
            table td.label {{
                width: 25%;
                text-align: center;
                font-weight: bold;
                background-color: #f9f9f9;
            }}
            table td.value {{
                width: 75%;
                text-align: left;
                padding-left: 15px;
            }}
            .signature {{
                text-align: center;
                margin-top: 50px;
            }}
            .date {{
                font-size: 13pt;
                margin-bottom: 40px;
            }}
            .company-signature {{
                font-size: 18pt;
                font-weight: bold;
                margin-top: 10px;
                position: relative;
                display: inline-block;
            }}
            .ceo-text {{
                display: inline-block;
                font-size: 13pt;
                margin-left: 10px;
                position: relative;
            }}
            .stamp {{
                position: absolute;
                right: -5px;
                top: -15px;
                z-index: 10;
            }}
            .stamp img {{
                width: 50px;
                height: 50px;
            }}
            .in-text {{
                position: relative;
                display: inline-block;
            }}
            @media print {{
                body {{
                    padding: 0;
                }}
                .container {{
                    max-width: 100%;
                }}
            }}
        </style>
    </head>
    <body>
        <div class="container">
            <div class="watermark">
                <img src="data:image/png;base64,{_get_base64_image('graylogo.png')}" alt="워터마크">
            </div>
            <div class="content">
                <div class="title">퇴&nbsp;직&nbsp;금&nbsp;정&nbsp;산&nbsp;서</div>
                <div class="title-underline"></div>
                <br><br>
                <table>
                    <tr>
                        <td class="label">성&nbsp;&nbsp;&nbsp;&nbsp;명</td>
                        <td class="value">{employee_data.get('name', '')}</td>
                    </tr>
                    <tr>
                        <td class="label">주민등록번호</td>
                        <td class="value">{employee_data.get('ssn', '') or ''}</td>
                    </tr>
                    <tr>
                        <td class="label">생년월일</td>
                        <td class="value">{employee_data.get('birth_date', '') or ''}</td>
                    </tr>
                    <tr>
                        <td class="label">소&nbsp;&nbsp;&nbsp;&nbsp;속</td>
                        <td class="value">도원반점 검단점</td>
                    </tr>
                    <tr>
                        <td class="label">입 사 일</td>
                        <td class="value">{employee_data.get('hire_date', '') or ''}</td>
                    </tr>
                    <tr>
                        <td class="label">퇴 사 일</td>
                        <td class="value">{resign_date}</td>
                    </tr>
                </table>
                <table>
                    <tr>
                        <td class="label">퇴직금</td>
                        <td class="value">{severance_amount:,}원</td>
                    </tr>
                </table>
                <div class="signature">
                    <div class="date">{today}</div>
                    <div class="company-signature">
                        도원반점 검단점
                        <span class="ceo-text">사장 김서은 (<span class="in-text">인<span class="stamp">
                            <img src="data:image/png;base64,{_get_base64_image('stamp.png')}" alt="직인">
                        </span></span>)</span>
                    </div>
                </div>
            </div>
        </div>
    </body>
    </html>
    """
    return html


# 서류 타입 매핑 (HTML) - 함수 정의 이후에 위치
DOCUMENT_GENERATORS_HTML = {
    'receipt_of_employment': create_receipt_of_employment_html,
    'career_certificate': create_career_certificate_html,
    'pay_stub': create_pay_stub_html,
    'withholding_receipt': create_withholding_receipt_html,
    'resignation_certificate': create_resignation_certificate_html,
    'severance_settlement': create_severance_settlement_html,
}


def create_receipt_of_employment(employee_data: dict) -> Document:
    """재직증명서 생성"""
    doc = Document()
    
    # 페이지 설정 (A4)
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    
    # 워터마크 추가
    add_watermark_to_section(section)
    
    # 기본 스타일 설정
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(11)
    
    # 제목 (큰 글씨, 중앙 정렬)
    title_para = doc.add_paragraph()
    title_run = title_para.add_run('재   직   증   명   서')
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')
    doc.add_paragraph('')
    
    # 본문
    body_para = doc.add_paragraph('위 사람은 당 업체에 재직 중임을 증명합니다.')
    body_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    body_para.runs[0].font.size = Pt(12)
    doc.add_paragraph('')
    doc.add_paragraph('')
    
    # 표 생성
    table = doc.add_table(rows=7, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(1.5)
    table.columns[1].width = Inches(4.5)
    
    rows_data = [
        ("성    명", employee_data.get('name', '')),
        ("생년월일", employee_data.get('birth_date', '') or ''),
        ("주    소", employee_data.get('address', '') or ''),
        ("소    속", "도원반점 검단점"),
        ("직    위", _format_position_title(employee_data)),
        ("입 사 일", employee_data.get('hire_date', '') or ''),
        ("퇴 사 일", employee_data.get('resign_date', '') or ''),
    ]
    
    for i, (label, value) in enumerate(rows_data):
        cell0 = table.rows[i].cells[0]
        cell1 = table.rows[i].cells[1]
        
        # 라벨 셀
        label_para = cell0.paragraphs[0]
        label_run = label_para.add_run(label)
        label_run.font.bold = True
        label_run.font.size = Pt(11)
        label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 값 셀
        value_para = cell1.paragraphs[0]
        value_run = value_para.add_run(value if value else '')
        value_run.font.size = Pt(11)
        value_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # 행 높이 설정 (750 twips = 약 1.32cm)
        tr = table.rows[i]._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), '750')
        trHeight.set(qn('w:hRule'), 'atLeast')
        trPr.append(trHeight)
    
    doc.add_paragraph('')
    doc.add_paragraph('')
    
    # 발급일 (중앙 정렬)
    today = datetime.now().strftime('%Y년 %m월 %d일')
    date_para = doc.add_paragraph(today)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.size = Pt(12)
    doc.add_paragraph('')
    
    # 발급자 (중앙 정렬, 굵게)
    company_para = doc.add_paragraph('도원반점 검단점')
    company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    company_para.runs[0].font.size = Pt(16)
    company_para.runs[0].font.bold = True
    
    ceo_para = doc.add_paragraph()
    ceo_run = ceo_para.add_run('사장 김서은 (인)')
    ceo_run.font.size = Pt(12)
    ceo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 직인 이미지 추가
    add_stamp_image(ceo_para)
    
    return doc


def create_career_certificate(employee_data: dict) -> Document:
    """경력증명서 생성"""
    doc = Document()
    
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    
    # 워터마크 추가
    add_watermark_to_section(section)
    
    # 기본 스타일 설정
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(11)
    
    # 제목
    title_para = doc.add_paragraph()
    title_run = title_para.add_run('경   력   증   명   서')
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')
    doc.add_paragraph('')
    
    # 본문
    body_para = doc.add_paragraph('위 사람은 본 식당에 근무한 경력이 아래와 같음을 증명합니다.')
    body_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    body_para.runs[0].font.size = Pt(12)
    doc.add_paragraph('')
    
    # 표 생성
    table = doc.add_table(rows=9, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(1.5)
    table.columns[1].width = Inches(4.5)
    
    rows_data = [
        ("성    명", employee_data.get('name', '')),
        ("생년월일", employee_data.get('birth_date', '') or ''),
        ("주    소", employee_data.get('address', '') or ''),
        ("소    속", "도원반점 검단점"),
        ("직    위", _format_position_title(employee_data)),
        ("입 사 일", employee_data.get('hire_date', '') or ''),
        ("퇴 사 일", employee_data.get('resign_date', '') or '재직중'),
        ("담당업무", _format_job_description(employee_data)),
        ("경    력", employee_data.get('career_period', '') or ''),
    ]
    
    for i, (label, value) in enumerate(rows_data):
        cell0 = table.rows[i].cells[0]
        cell1 = table.rows[i].cells[1]
        
        # 라벨 셀
        label_para = cell0.paragraphs[0]
        label_run = label_para.add_run(label)
        label_run.font.bold = True
        label_run.font.size = Pt(11)
        label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 값 셀
        value_para = cell1.paragraphs[0]
        value_run = value_para.add_run(value if value else '')
        value_run.font.size = Pt(11)
        value_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # 행 높이 설정 (580 twips)
        tr = table.rows[i]._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), '580')
        trHeight.set(qn('w:hRule'), 'atLeast')
        trPr.append(trHeight)
    
    doc.add_paragraph('')
    doc.add_paragraph('')
    
    # 발급일
    today = datetime.now().strftime('%Y년 %m월 %d일')
    date_para = doc.add_paragraph(today)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.size = Pt(12)
    doc.add_paragraph('')
    
    # 발급자
    company_para = doc.add_paragraph('도원반점 검단점')
    company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    company_para.runs[0].font.size = Pt(16)
    company_para.runs[0].font.bold = True
    
    ceo_para = doc.add_paragraph()
    ceo_run = ceo_para.add_run('사장 김서은 (인)')
    ceo_run.font.size = Pt(12)
    ceo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 직인 이미지 추가
    add_stamp_image(ceo_para)
    
    return doc


def create_pay_stub(employee_data: dict, payroll_data: Optional[dict] = None) -> Document:
    """급여명세서 생성"""
    doc = Document()
    
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    
    # 워터마크 추가
    add_watermark_to_section(section)
    
    title = doc.add_heading('급여명세서', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 기본 정보 표
    info_table = doc.add_table(rows=3, cols=2)
    try:
        info_table.style = 'Light Grid Accent 1'
    except:
        pass
    
    today = datetime.now()
    year_month = payroll_data.get('year_month', f'{today.year}-{today.month:02d}') if payroll_data else f'{today.year}-{today.month:02d}'
    
    info_rows = [
        ("성    명", employee_data.get('name', '')),
        ("소    속", "도원반점 검단점"),
        ("지 급 년 월", year_month),
    ]
    
    for i, (label, value) in enumerate(info_rows):
        cell0 = info_table.rows[i].cells[0]
        cell1 = info_table.rows[i].cells[1]
        cell0.text = label
        cell1.text = value
        if len(cell0.paragraphs) > 0 and len(cell0.paragraphs[0].runs) > 0:
            cell0.paragraphs[0].runs[0].font.bold = True
        cell0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('')
    
    # 급여 내역 표
    detail_table = doc.add_table(rows=5, cols=2)
    try:
        detail_table.style = 'Light Grid Accent 1'
    except:
        pass
    
    base_pay = payroll_data.get('base_pay', 0) if payroll_data else 0
    deductions = payroll_data.get('deductions', 0) if payroll_data else 0
    net_pay = payroll_data.get('net_pay', 0) if payroll_data else 0
    
    detail_rows = [
        ("기본급", f"{int(base_pay):,}원"),
        ("공제액", f"{int(deductions):,}원"),
        ("실수령액", f"{int(net_pay):,}원"),
    ]
    
    for i, (label, value) in enumerate(detail_rows):
        cell0 = detail_table.rows[i].cells[0]
        cell1 = detail_table.rows[i].cells[1]
        cell0.text = label
        cell1.text = value
        if len(cell0.paragraphs) > 0 and len(cell0.paragraphs[0].runs) > 0:
            cell0.paragraphs[0].runs[0].font.bold = True
        cell0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph(f'{datetime.now().strftime("%Y년 %m월 %d일")}')
    doc.add_paragraph('도원반점 검단점')
    
    ceo_para = doc.add_paragraph()
    ceo_run = ceo_para.add_run('사장 김서은 (인)')
    ceo_run.font.size = Pt(12)
    ceo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 직인 이미지 추가
    add_stamp_image(ceo_para)
    
    return doc


def create_withholding_receipt(employee_data: dict, payroll_data: Optional[dict] = None) -> Document:
    """원천징수영수증 생성"""
    doc = Document()
    
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    
    # 워터마크 추가
    add_watermark_to_section(section)
    
    title = doc.add_heading('원천징수영수증', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 기본 정보
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Light Grid Accent 1'
    
    today = datetime.now()
    year_month = payroll_data.get('year_month', f'{today.year}-{today.month:02d}') if payroll_data else f'{today.year}-{today.month:02d}'
    
    info_rows = [
        ("성    명", employee_data.get('name', '')),
        ("주민등록번호", employee_data.get('birth_date', '') or ''),
        ("소    속", "도원반점 검단점"),
        ("지 급 년 월", year_month),
    ]
    
    for i, (label, value) in enumerate(info_rows):
        cell0 = info_table.rows[i].cells[0]
        cell1 = info_table.rows[i].cells[1]
        cell0.text = label
        cell1.text = value
        if len(cell0.paragraphs) > 0 and len(cell0.paragraphs[0].runs) > 0:
            cell0.paragraphs[0].runs[0].font.bold = True
        cell0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('')
    
    # 징수 내역
    detail_table = doc.add_table(rows=3, cols=2)
    try:
        detail_table.style = 'Light Grid Accent 1'
    except:
        pass
    
    base_pay = payroll_data.get('base_pay', 0) if payroll_data else 0
    deductions = payroll_data.get('deductions', 0) if payroll_data else 0
    
    detail_rows = [
        ("지 급 금 액", f"{int(base_pay):,}원"),
        ("원천징수세액", f"{int(deductions):,}원"),
    ]
    
    for i, (label, value) in enumerate(detail_rows):
        cell0 = detail_table.rows[i].cells[0]
        cell1 = detail_table.rows[i].cells[1]
        cell0.text = label
        cell1.text = value
        if len(cell0.paragraphs) > 0 and len(cell0.paragraphs[0].runs) > 0:
            cell0.paragraphs[0].runs[0].font.bold = True
        cell0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph(f'{datetime.now().strftime("%Y년 %m월 %d일")}')
    doc.add_paragraph('도원반점 검단점')
    
    ceo_para = doc.add_paragraph()
    ceo_run = ceo_para.add_run('사장 김서은 (인)')
    ceo_run.font.size = Pt(12)
    ceo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 직인 이미지 추가
    add_stamp_image(ceo_para)
    
    return doc


def create_resignation_certificate(employee_data: dict) -> Document:
    """퇴직증명서 생성"""
    doc = Document()
    
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    
    # 워터마크 추가
    add_watermark_to_section(section)
    
    title = doc.add_heading('퇴직증명서', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('위 사람은 본 식당의 직원으로 재직하다가 아래와 같이 퇴직함을 증명합니다.')
    doc.add_paragraph('')
    
    table = doc.add_table(rows=8, cols=2)
    try:
        table.style = 'Light Grid Accent 1'
    except:
        pass
    
    rows_data = [
        ("성    명", employee_data.get('name', '')),
        ("생년월일", employee_data.get('birth_date', '') or ''),
        ("주    소", employee_data.get('address', '') or ''),
        ("소    속", "도원반점 검단점"),
        ("직    위", _format_position_title(employee_data)),
        ("입 사 일", employee_data.get('hire_date', '') or ''),
        ("퇴 사 일", employee_data.get('resign_date', '') or datetime.now().strftime('%Y-%m-%d')),
        ("담당업무", _format_job_description(employee_data)),
    ]
    
    for i, (label, value) in enumerate(rows_data):
        cell0 = table.rows[i].cells[0]
        cell1 = table.rows[i].cells[1]
        cell0.text = label
        cell1.text = value
        # Bold 설정
        if len(cell0.paragraphs) > 0 and len(cell0.paragraphs[0].runs) > 0:
            cell0.paragraphs[0].runs[0].font.bold = True
        cell0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph(f'{datetime.now().strftime("%Y년 %m월 %d일")}')
    doc.add_paragraph('도원반점 검단점')
    
    ceo_para = doc.add_paragraph()
    ceo_run = ceo_para.add_run('사장 김서은 (인)')
    ceo_run.font.size = Pt(12)
    ceo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 직인 이미지 추가
    add_stamp_image(ceo_para)
    
    return doc


def create_severance_settlement(employee_data: dict, severance_data: Optional[dict] = None) -> Document:
    """퇴직금정산서 생성"""
    doc = Document()
    
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    
    # 워터마크 추가
    add_watermark_to_section(section)
    
    title = doc.add_heading('퇴직금정산서', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 기본 정보
    info_table = doc.add_table(rows=5, cols=2)
    try:
        info_table.style = 'Light Grid Accent 1'
    except:
        pass
    
    info_rows = [
        ("성    명", employee_data.get('name', '')),
        ("생년월일", employee_data.get('birth_date', '') or ''),
        ("소    속", "도원반점 검단점"),
        ("입 사 일", employee_data.get('hire_date', '') or ''),
        ("퇴 사 일", employee_data.get('resign_date', '') or datetime.now().strftime('%Y-%m-%d')),
    ]
    
    for i, (label, value) in enumerate(info_rows):
        cell0 = info_table.rows[i].cells[0]
        cell1 = info_table.rows[i].cells[1]
        cell0.text = label
        cell1.text = value
        if len(cell0.paragraphs) > 0 and len(cell0.paragraphs[0].runs) > 0:
            cell0.paragraphs[0].runs[0].font.bold = True
        cell0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('')
    
    # 정산 내역
    detail_table = doc.add_table(rows=3, cols=2)
    try:
        detail_table.style = 'Light Grid Accent 1'
    except:
        pass
    
    severance_amount = severance_data.get('severance_amount', 0) if severance_data else 0
    
    detail_rows = [
        ("퇴직금", f"{int(severance_amount):,}원"),
    ]
    
    for i, (label, value) in enumerate(detail_rows):
        cell0 = detail_table.rows[i].cells[0]
        cell1 = detail_table.rows[i].cells[1]
        cell0.text = label
        cell1.text = value
        if len(cell0.paragraphs) > 0 and len(cell0.paragraphs[0].runs) > 0:
            cell0.paragraphs[0].runs[0].font.bold = True
        cell0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph(f'{datetime.now().strftime("%Y년 %m월 %d일")}')
    doc.add_paragraph('도원반점 검단점')
    
    ceo_para = doc.add_paragraph()
    ceo_run = ceo_para.add_run('사장 김서은 (인)')
    ceo_run.font.size = Pt(12)
    ceo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 직인 이미지 추가
    add_stamp_image(ceo_para)
    
    return doc


# 서류 타입 매핑
DOCUMENT_GENERATORS = {
    'receipt_of_employment': create_receipt_of_employment,
    'career_certificate': create_career_certificate,
    'pay_stub': create_pay_stub,
    'withholding_receipt': create_withholding_receipt,
    'resignation_certificate': create_resignation_certificate,
    'severance_settlement': create_severance_settlement,
}


def generate_document_html(doc_type: str, employee_data: dict, additional_data: Optional[dict] = None) -> str:
    """서류 생성 메인 함수 (HTML 형식)"""
    generator = DOCUMENT_GENERATORS_HTML.get(doc_type)
    if not generator:
        raise ValueError(f"Unknown document type: {doc_type}")
    
    if doc_type in ['pay_stub', 'withholding_receipt']:
        return generator(employee_data, additional_data)
    elif doc_type == 'severance_settlement':
        return generator(employee_data, additional_data)
    else:
        return generator(employee_data)


def generate_document(doc_type: str, employee_data: dict, additional_data: Optional[dict] = None) -> Document:
    """서류 생성 메인 함수 (DOCX 형식 - 기존)"""
    generator = DOCUMENT_GENERATORS.get(doc_type)
    if not generator:
        raise ValueError(f"Unknown document type: {doc_type}")
    
    if doc_type in ['pay_stub', 'withholding_receipt']:
        return generator(employee_data, additional_data)
    elif doc_type == 'severance_settlement':
        return generator(employee_data, additional_data)
    else:
        return generator(employee_data)

